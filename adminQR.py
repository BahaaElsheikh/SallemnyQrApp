import cv2
import sqlite3
from pyzbar.pyzbar import decode
from kivy.app import App
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.uix.scrollview import ScrollView
from kivy.uix.gridlayout import GridLayout
from kivy.uix.image import Image
from kivy.clock import Clock
from kivy.core.window import Window
from openpyxl import Workbook
from docx import Document
from threading import Thread
from kivy.graphics.texture import Texture
import datetime
from kivy.core.image import Image as CoreImage
from kivy.uix.textinput import TextInput
from kivy.uix.popup import Popup
import qrcode
from kivy.uix.screenmanager import Screen
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.anchorlayout import AnchorLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
import os

# ====   Screen Settings  ====
Window.clearcolor = (1, 1, 1, 1)  #  White BackGround
Window.size = (365, 640)  #



# ========== Data Base ==========
def save_to_db(data,st):
    conn = sqlite3.connect("qr_data.db")
    c = conn.cursor()
    c.execute("CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY, info TEXT,scan_time TEXT)")
    # c.execute("ALTER TABLE users ADD COLUMN scan_time TEXT")
    c.execute("INSERT INTO users (info,scan_time) VALUES (?,?)", (data,st))
    conn.commit()
    conn.close()

def fetch_all_data():
    conn = sqlite3.connect("qr_data.db")
    c = conn.cursor()
    c.execute("SELECT * FROM users")
    data = c.fetchall()
    conn.close()
    return data

def delete_all_data():
    conn = sqlite3.connect("qr_data.db")
    c = conn.cursor()
    c.execute("DELETE FROM users")
    conn.commit()
    conn.close()

            # text=" Warning!! \n File Name Shouldn't Contain Any of These Characters: \\ / : * ? \" < > | \n Enter File Name Only Not the Extension ",


# ========== Export Word ==========
def export_to_word(data,Fname="qr_data"):
    doc = Document()
    doc.add_heading("Scanned QR Data", 0)
    for item in data:
        doc.add_paragraph(f"ID: {item[0]}, Info: {item[1]} , Time: {item[2]}")
    doc.save(f"{Fname}.docx")

# ========== Export Excel ==========
def export_to_excel(data,Fname="qr_data"):
    wb = Workbook()
    ws = wb.active
    ws.append(["ID", "Info"])
    for item in data:
        ws.append([item[0], item[1]])
    wb.save(f"{Fname}.xlsx")

# def store_Data(myData):
#     return myData


class ExportOption(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        
        main_layout = BoxLayout(orientation='vertical', spacing=15, padding=20)

     
        title = Label(
            text="Export",
            font_size='40sp',
            bold=True,
            color=(0, 0, 0, 1),
            size_hint=(1, None),
            height=40
        )
        main_layout.add_widget(title)

        # Anchor layout To MAke it in the center of the Pageا
        center_anchor = AnchorLayout(anchor_x='center', anchor_y='center', size_hint=(1, None), height=160)

        center_layout = BoxLayout(orientation='vertical', spacing=10, size_hint=(None, None), width=300, height=160)

        # Text input لحقل الاسم
        self.file_input = TextInput(
            hint_text="Enter File Name",
            multiline=False,
            size_hint=(1, None),
            height=60,
            font_size='20sp',
            background_color=(1, 1, 1, 1),
            foreground_color=(0, 0, 0, 1),
            padding=(10, 10)
        )

        # تحذير
        self.instructions = Label(
            text='Avoid using: \\ / : * ? " < . > | ',
            font_size='12sp',
            color=(1, 0, 0, 1),
            size_hint=(1, None),
            height=20
        )

        center_layout.add_widget(self.file_input)
        center_layout.add_widget(self.instructions)
        center_anchor.add_widget(center_layout)

        main_layout.add_widget(center_anchor)

        # الأزرار
        export_word_btn = Button(
            text="Export to Word",
            color=(1, 1, 1, 1),
            background_color=(0.0, 0.0, 0.95, 1),
            size_hint=(1, None),
            height=80
        )
        export_word_btn.bind(on_press=self.export_word)

        export_excel_btn = Button(
            text="Export to Excel",
            color=(1, 1, 1, 1),
            background_color=(0.0, 1, 0, 1),
            size_hint=(1, None),
            height=80
        )
        export_excel_btn.bind(on_press=self.export_excel)

        back_btn = Button(
            text="Back to Menu",
            color=(1, 1, 1, 1),
            background_color=(0.5, 0.5, 0.5, 0.5),
            size_hint=(1, None),
            height=80
        )
        back_btn.bind(on_press=lambda x: setattr(self.manager, 'current', 'menu'))

        main_layout.add_widget(export_word_btn)
        main_layout.add_widget(export_excel_btn)
        main_layout.add_widget(back_btn)

        self.add_widget(main_layout)

    def export_word(self, instance):
        file_name = self.file_input.text.strip()
        print(f"Exporting to Word: {file_name}")
        export_to_word(fetch_all_data(), file_name)

    def export_excel(self, instance):
        file_name = self.file_input.text.strip()
        print(f"Exporting to Excel: {file_name}")
        export_to_excel(fetch_all_data(), file_name)





# ========== Splash Screen ==========
class SplashScreen(Screen):
    def on_enter(self):
        layout = BoxLayout(orientation='vertical')
        layout.add_widget(Image(source='LOGO.png', allow_stretch=True, keep_ratio=True))
        self.add_widget(layout)
        Clock.schedule_once(self.switch_to_main, 2)

    def switch_to_main(self, dt):
        self.manager.current = 'menu'

# ========== Main Menu ==========
class MainMenu(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        layout = BoxLayout(orientation="vertical", spacing=20, padding=20)

        logo = Image(source="LOGO.png")
        layout.add_widget(logo)

        create_btn = Button(
            text="Create QR Code",
            size_hint=(0.6, 0.1),
            pos_hint={'center_x': 0.5},
            bold=True, 
            background_color=(0.5, 0.7, 1, .9),
            color=(1, 1, 1, 1)
        )

        create_btn.bind(on_press=lambda x: setattr(self.manager, 'current', 'create'))


        scan_btn = Button(
            text="Scan QR Code",
            size_hint=(0.6, 0.1),
            pos_hint={'center_x': 0.5},
            bold=True, 
            background_color=(0.5, 0.7, 1, .9),
            color=(1, 1, 1, 1)
        )
        scan_btn.bind(on_press=self.go_to_scan)

        show_btn = Button(
            text="Show Data",
            size_hint=(0.6, 0.1),
            pos_hint={'center_x': 0.5},
            background_color=(0.5, 0.7, 1, .9),
            bold=True, 
            color=(1, 1, 1, 1)
        )
        show_btn.bind(on_press=lambda x: self.manager.get_screen('data').update_data())
        show_btn.bind(on_press=lambda x: setattr(self.manager, 'current', 'data'))


        layout.add_widget(create_btn)
        layout.add_widget(scan_btn)
        layout.add_widget(show_btn)

        self.add_widget(layout)

    def go_to_scan(self, instance):
        self.manager.get_screen('scan').start_scan()
        self.manager.current = 'scan'

# ========== Scan Screen ==========
class ScanScreen(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.layout = BoxLayout(orientation="vertical")
        self.image = Image()
        self.layout.add_widget(self.image)

        back_btn = Button(
            text="Back to Menu",
            size_hint=(1, 0.1),
            background_color=(0.5, 0.7, 1, 1),
            color=(1, 1, 1, 1)
        )
        back_btn.bind(on_press=self.stop_scanning)
        self.layout.add_widget(back_btn)

        self.add_widget(self.layout)

        self.cap = None
        self.running = False
        Clock.schedule_interval(self.update_frame, 1.0 / 30.0)

    def start_scan(self):
        self.cap = cv2.VideoCapture(0)
        self.running = True

    def stop_scanning(self, *args):
        self.running = False
        if self.cap:
            self.cap.release()
        self.manager.current = 'menu'

    def update_frame(self, dt):
        if self.running and self.cap and self.cap.isOpened():
            ret, frame = self.cap.read()
            if ret:
                buf1 = cv2.flip(frame, 0)
                buf = buf1.tobytes()
                image_texture = Texture.create(size=(frame.shape[1], frame.shape[0]), colorfmt='bgr')
                image_texture.blit_buffer(buf, colorfmt='bgr', bufferfmt='ubyte')
                self.image.texture = image_texture

                for code in decode(frame):
                    data = code.data.decode('utf-8')
                    scan_time =datetime.datetime.now().strftime('%Y-%m-%d %H:%M')  # الوقت بشكل مقروء

                    save_to_db(data,scan_time)
                    self.running = False
                    if self.cap:
                        self.cap.release()
                    # Clock.schedule_once(lambda dt: setattr(self.manager, 'current', 'menu'))
                    Clock.schedule_once(lambda dt: setattr(self.manager, 'current', 'confirm'))
                    global myData 
                    myData = data
                    break

#==========Confirm screen ==========
class Confirm(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        layout = BoxLayout(orientation="vertical", spacing=20, padding=20)
        conf = Image(
            source="Confirm.png",
            size_hint=(1, 1),  
            allow_stretch=True,
            keep_ratio=False,
            pos_hint={'center_x': 0.5, 'top': 0.95},  

        )
        layout.add_widget(conf)

        confirm_label = Label(
            text=f"Confirm",
            font_size='70sp',          
            bold=True,                
            color=(0.09, 0.55, 0.22, 1),
            size_hint=(1, 0.9),        
            halign='center',           
            valign='middle'           
        )

        confirm_label.bind(size=confirm_label.setter('text_size'))  # Center The Text

        layout.add_widget(confirm_label)
        self.add_widget(layout)

    def on_enter(self):
     # أول ما الشاشة تفتح تبدأ العد التنازلي للرجوع
         Clock.schedule_once(self.switch_to_main,.45 )

    def switch_to_main(self, dt):
        self.manager.current = 'menu'
    


# ========== Data Screen ==========
class DataScreen(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.layout = BoxLayout(orientation="vertical", padding=10, spacing=10)
        self.scroll = ScrollView()
        self.data_grid = GridLayout(cols=1, size_hint_y=None, spacing=5)
        self.data_grid.bind(minimum_height=self.data_grid.setter('height'))
        self.scroll.add_widget(self.data_grid)
        self.layout.add_widget(self.scroll)

        btns = BoxLayout(orientation="vertical",size_hint=(1, 0.4), spacing=10)

        show_btn = Button(text="Export Data", color=(1, 1, 1, 1), background_color=(0.5, 0.7, 1, 1))
        # show_btn.bind(on_press=lambda x: self.manager.get_screen('export').update_data())
        show_btn.bind(on_press=lambda x: setattr(self.manager, 'current', 'export'))


        export_word_btn = Button(text="Export to Word", color=(0, 0, 0, 1), background_color=(0.5, 0.7, 1, 1))
        export_word_btn.bind(on_press=lambda x: export_to_word(fetch_all_data()))

        export_excel_btn = Button(text="Export to Excel", color=(0, 0, 0, 1), background_color=(0.5, 0.7, 1, 1))
        export_excel_btn.bind(on_press=lambda x: export_to_excel(fetch_all_data()))

        back_btn = Button(text=" Back to Menu", color=(1, 1, 1, 1), background_color=(0.5, 0.5, 0.5, 0.5))
        back_btn.bind(on_press=lambda x: setattr(self.manager, 'current', 'menu'))

        clear_btn = Button(text="  Clear Data ", color=(1, 1, 1, 1), background_color=(.95, 0.0, 0.0, .9))
        clear_btn.bind(on_press=self.clear_and_refresh_data)

        btns.add_widget(show_btn)
        # btns.add_widget(export_word_btn)
        # btns.add_widget(export_excel_btn)
        btns.add_widget(clear_btn)
        btns.add_widget(back_btn)

        self.layout.add_widget(btns)
        self.add_widget(self.layout)

    def update_data(self):
        
        self.data_grid.clear_widgets()
        data = fetch_all_data()
        if data ==[]:
           self.data_grid.add_widget(Label(text="No Data !!",size_hint_y=None,
                height=100,
                font_size='25sp',
                color=(1, 0, 0, 1),bold =True))

        for row in data:
            self.data_grid.add_widget(Label(
                text=f"\n\n\n  ID: {row[0]} | Info: {row[1]} | Time: {row[2]} \n {"_"*60} \n ",
                size_hint_y=None,
                height=40,
                font_size='11sp',
                color=(0, 0, 0, 1)
            ))

    def clear_and_refresh_data(self, instance):
        delete_all_data()
        self.update_data()


class CreateData(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        main_layout = BoxLayout(orientation='vertical', spacing=10, padding=20)

        title = Label(
            text="Enter Your Data ",
            font_size='30sp',
            bold=True,
            color=(.1, .3, 1, 1),
            size_hint=(1, None),
            height=40
        )
        main_layout.add_widget(title)

        
        center_anchor = AnchorLayout(anchor_x='center', anchor_y='center')

        center_layout = BoxLayout(orientation='vertical', spacing=10, size_hint=(None, None), size=(400, 300))

        self.name_input = TextInput(
            hint_text="Enter Your Name ",
            multiline=False,
            size_hint=(1, None),
            height=60,
            font_size='15sp',
            background_color=(1, 1, 1, 1),
            foreground_color=(0, 0, 0, 1),
            padding=(10, 10)
        )
        self.subj_input = TextInput(
            hint_text="Enter The Subject Name ",
            multiline=False,
            size_hint=(1, None),
            height=60,
            font_size='15sp',
            background_color=(1, 1, 1, 1),
            foreground_color=(0, 0, 0, 1),
            padding=(10, 10)
        )

        self.qr_btn = Button(
            text="Generate QR Code",
            size_hint=(1, None),
            background_color=(.3, .7, 1, 1),
            height=60,
            font_size='16sp'
        )
        self.qr_btn.bind(on_press=self.generate_qr)

        center_layout.add_widget(self.name_input)
        center_layout.add_widget(self.subj_input)
        center_layout.add_widget(self.qr_btn)

        center_anchor.add_widget(center_layout)
        main_layout.add_widget(center_anchor)

       
        bottom_anchor = AnchorLayout(anchor_x='center', anchor_y='bottom', size_hint=(1, 1))

        back_btn = Button(
            text="Back to Menu",
            size_hint=(1, None),
            height=40,
            font_size='16sp',
            background_color=(0.5, 0.5, 0.5, 0.5),
            color=(1, 1, 1, 1)
        )
        back_btn.bind(on_press=lambda x: setattr(self.manager, 'current', 'menu'))

        bottom_anchor.add_widget(back_btn)
        main_layout.add_widget(bottom_anchor)

        self.add_widget(main_layout)

    def generate_qr(self, instance):
        Name = self.name_input.text.strip()
        subject = self.subj_input.text.strip()
        data = f"{Name} | {subject}"

        if  data.strip()==" | ":
            popup = Popup(title="Warning !!", content=Label(text=" Empty Data"), size_hint=(0.6, 0.3))
            popup.open()
            return
        create_time =datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
        filename = f"my_qr_{create_time}.png"
        qr = qrcode.make(data)
        qr.save(filename)

        self.manager.get_screen('qr').qr_image.source = ""
        self.manager.get_screen('qr').qr_image.source = filename

        self.manager.current = 'qr'

        self.last_qr_file = filename  ## Saving Last File Name 



class QRScreen(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        layout = BoxLayout(orientation="vertical", padding=20, spacing=20)

        self.qr_image = Image(source="", size_hint=(1, 0.8))
        layout.add_widget(self.qr_image)

        self.back_btn = Button(text="Back to Create ", size_hint=(1, 0.2))
        self.back_btn.bind(on_press=self.go_back)       
            
        layout.add_widget(self.back_btn)    
        self.add_widget(layout)

    def go_back(self, instance):

        ##Delete The Previous QR code File
        create_screen = self.manager.get_screen('create')
        if hasattr(create_screen, 'last_qr_file'):
            try:
                os.remove(create_screen.last_qr_file)
            except Exception as e:
                print(f"Error deleting file: {e}")
        
        self.manager.current = 'create'



# ========== Main App ==========
class SallemnyApp(App):

    def build(self):
        sm = ScreenManager()
        sm.add_widget(SplashScreen(name='splash'))
        sm.add_widget(MainMenu(name='menu'))
        sm.add_widget(ScanScreen(name='scan'))
        sm.add_widget(DataScreen(name='data'))
        sm.add_widget(ExportOption(name='export'))
        sm.add_widget(Confirm(name='confirm'))
        sm.add_widget(CreateData(name='create'))
        sm.add_widget(QRScreen(name='qr'))
        sm.current = 'splash'
        return sm

if __name__ == "__main__":
    SallemnyApp().run()
